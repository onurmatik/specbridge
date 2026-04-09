import io
import secrets
from html import escape
from os.path import splitext

from django.urls import reverse
from django.utils import timezone
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from exports.models import ExportArtifact, ExportFileType, ExportFormat
from specs.models import AuditEventType
from specs.services import (
    log_audit_event,
    section_summaries,
    section_title_for_ref,
)
from specs.spec_document import markdown_to_blocks


def selected_sections_for_export(project, configuration: dict | None = None):
    configuration = configuration or {}
    section_ids = configuration.get("section_ids", "")
    if isinstance(section_ids, str):
        allowed = {section_id.strip() for section_id in section_ids.split(",") if section_id.strip()}
    elif isinstance(section_ids, list):
        allowed = {section_id for section_id in section_ids if section_id}
    else:
        allowed = set()

    sections = list(section_summaries(project))
    if not allowed:
        return sections
    return [section for section in sections if section["id"] in allowed]


def normalize_export_format(export_format: str) -> str:
    normalized = f"{export_format or ''}".strip().lower()
    if normalized not in ExportFormat.values:
        raise ValueError(f"Unsupported export format: {export_format}")
    return normalized


def _normalize_file_type_value(value: str | None) -> str | None:
    normalized = f"{value or ''}".strip().lower().lstrip(".")
    if not normalized:
        return None
    aliases = {
        "markdown": ExportFileType.MARKDOWN,
        "md": ExportFileType.MARKDOWN,
        "pdf": ExportFileType.PDF,
        "docx": ExportFileType.DOCX,
    }
    return aliases.get(normalized)


def normalize_requested_export_file_type(configuration: dict | None = None) -> str:
    configuration = configuration or {}
    raw_file_type = configuration.get("file_type")
    requested = _normalize_file_type_value(raw_file_type)
    if requested:
        return requested
    if raw_file_type not in (None, ""):
        raise ValueError(f"Unsupported export file type: {raw_file_type}")

    extension = configuration.get("extension")
    normalized_extension = _normalize_file_type_value(extension)
    if normalized_extension:
        return normalized_extension

    if extension not in (None, ""):
        raise ValueError(f"Unsupported export file type: {extension}")
    return ExportFileType.MARKDOWN


def export_file_type_for_artifact(artifact: ExportArtifact) -> str:
    configuration = artifact.configuration or {}
    configured_type = _normalize_file_type_value(configuration.get("file_type"))
    if configured_type:
        return configured_type

    extension_type = _normalize_file_type_value(configuration.get("extension"))
    if extension_type:
        return extension_type

    filename_type = _normalize_file_type_value(splitext(artifact.filename or "")[1])
    if filename_type:
        return filename_type
    return ExportFileType.MARKDOWN


def export_content_type(file_type: str) -> str:
    return {
        ExportFileType.MARKDOWN: "text/markdown; charset=utf-8",
        ExportFileType.PDF: "application/pdf",
        ExportFileType.DOCX: "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }[file_type]


def download_url_for_artifact(artifact: ExportArtifact) -> str:
    return reverse("project-export-download", args=[artifact.project.slug, artifact.id])


def build_export_content(project, export_format: str, configuration: dict | None = None) -> str:
    configuration = configuration or {}
    resolved_format = normalize_export_format(export_format)
    include_resolved_questions = configuration.get("include_resolved_questions", False)
    sections = selected_sections_for_export(project, configuration)
    allowed_ids = {section["id"] for section in sections}
    lines = [
        f"# {project.name}",
        "",
        project.summary,
        "",
        "## Sections",
    ]
    for index, section in enumerate(sections, start=1):
        lines.extend(
            [
                f"### {index}. {section['title']}",
                f"- Status: {section['status'].replace('-', ' ').title()}",
                f"- Type: {section['kind']}",
                "",
                section["body"] or "_No content yet._",
                "",
            ]
        )
    lines.append("## Decisions")
    for decision in project.decisions.all():
        primary_ref = decision.primary_ref or {}
        if primary_ref.get("section_id") and primary_ref["section_id"] not in allowed_ids:
            continue
        related_label = f" ({section_title_for_ref(project, primary_ref)})" if primary_ref else ""
        lines.append(f"- [{decision.status}] {decision.title}{related_label}: {decision.summary}")
    lines.append("")
    lines.append("## Assumptions")
    for assumption in project.assumptions.all():
        primary_ref = assumption.primary_ref or {}
        if primary_ref.get("section_id") and primary_ref["section_id"] not in allowed_ids:
            continue
        related_label = f" ({section_title_for_ref(project, primary_ref)})" if primary_ref else ""
        lines.append(f"- [{assumption.status}] {assumption.title}{related_label}: {assumption.description}")
    if include_resolved_questions:
        lines.append("")
        lines.append("## Questions")
        for question in project.questions.all():
            primary_ref = question.primary_ref or {}
            if primary_ref.get("section_id") and primary_ref["section_id"] not in allowed_ids:
                continue
            related_label = f" ({section_title_for_ref(project, primary_ref)})" if primary_ref else ""
            lines.append(f"- [{question.status}] {question.title}{related_label}")

    if resolved_format == ExportFormat.AGENT:
        lines = [
            "You are implementing a single spec document workspace in SpecBridge.",
            "",
            *lines,
        ]
    elif resolved_format == ExportFormat.UI_UX_AGENT:
        lines = [
            "Based on the spec below, write a prompt to the UI/UX agent to design the defined app and workflows.",
            "",
            "Treat the spec as the source of truth. The prompt should cover:",
            "- the key screens and navigation structure",
            "- the primary user flows and workflow transitions",
            "- empty, loading, success, and error states",
            "- edge cases, interaction details, and handoff considerations",
            "",
            *lines,
        ]

    return "\n".join(lines)


def create_export(project, export_format, actor, configuration=None):
    configuration = configuration or {}
    resolved_format = normalize_export_format(export_format)
    file_type = normalize_requested_export_file_type(configuration)
    normalized_configuration = {**configuration, "file_type": file_type}
    label = dict(ExportFormat.choices).get(resolved_format, resolved_format)
    stamp = timezone.localtime().strftime("%Y%m%d-%H%M")
    artifact = ExportArtifact.objects.create(
        project=project,
        format=resolved_format,
        title=f"{label} export",
        filename=f"{project.slug}_{resolved_format}_{stamp}.{file_type}",
        generated_by=actor,
        configuration=normalized_configuration,
        content=build_export_content(project, resolved_format, normalized_configuration),
        share_enabled=configuration.get("share_enabled", False),
        share_token=secrets.token_hex(12),
    )
    log_audit_event(
        project=project,
        actor=actor,
        event_type=AuditEventType.EXPORT_CREATED,
        title=f"Generated export: {artifact.filename}",
        description=f"{label} export generated",
        metadata={"export_id": artifact.id, "format": artifact.format},
    )
    return artifact


def _inline_plain_text(nodes) -> str:
    parts: list[str] = []
    for node in nodes or []:
        if not isinstance(node, dict):
            continue
        if node.get("type") == "text":
            parts.append(node.get("text", ""))
            continue
        parts.append(_inline_plain_text(node.get("content", [])))
    return "".join(parts)


def _inline_reportlab_markup(nodes) -> str:
    parts: list[str] = []
    for node in nodes or []:
        if not isinstance(node, dict):
            continue
        if node.get("type") == "text":
            text = escape(node.get("text", "")).replace("\n", "<br/>")
            marks = {mark.get("type") for mark in node.get("marks", []) if isinstance(mark, dict)}
            if "bold" in marks:
                text = f"<b>{text}</b>"
            if "italic" in marks:
                text = f"<i>{text}</i>"
            parts.append(text)
            continue
        parts.append(_inline_reportlab_markup(node.get("content", [])))
    return "".join(parts)


def _append_docx_runs(paragraph, nodes) -> None:
    for node in nodes or []:
        if not isinstance(node, dict):
            continue
        if node.get("type") != "text":
            _append_docx_runs(paragraph, node.get("content", []))
            continue
        run = paragraph.add_run(node.get("text", ""))
        marks = {mark.get("type") for mark in node.get("marks", []) if isinstance(mark, dict)}
        run.bold = "bold" in marks
        run.italic = "italic" in marks


def _iter_renderable_elements(blocks, *, list_level: int = 0):
    for block in blocks or []:
        if not isinstance(block, dict):
            continue
        block_type = block.get("type")
        if block_type == "heading":
            yield {
                "kind": "heading",
                "level": max(int(block.get("attrs", {}).get("level", 1) or 1), 1),
                "content": block.get("content", []),
            }
            continue
        if block_type == "paragraph":
            if _inline_plain_text(block.get("content", [])).strip():
                yield {"kind": "paragraph", "level": list_level, "content": block.get("content", [])}
            continue
        if block_type == "bulletList":
            yield from _iter_list_elements(block.get("content", []), ordered=False, list_level=list_level)
            continue
        if block_type == "orderedList":
            start = max(int(block.get("attrs", {}).get("start", 1) or 1), 1)
            yield from _iter_list_elements(block.get("content", []), ordered=True, list_level=list_level, start=start)


def _iter_list_elements(items, *, ordered: bool, list_level: int, start: int = 1):
    index = start
    for item in items or []:
        if not isinstance(item, dict):
            continue
        children = item.get("content", [])
        paragraph = next(
            (child for child in children if isinstance(child, dict) and child.get("type") == "paragraph"),
            None,
        )
        if paragraph and _inline_plain_text(paragraph.get("content", [])).strip():
            yield {
                "kind": "list_item",
                "level": list_level,
                "ordered": ordered,
                "index": index,
                "content": paragraph.get("content", []),
            }
        for child in children:
            if not isinstance(child, dict):
                continue
            child_type = child.get("type")
            if child_type == "paragraph" and child is not paragraph:
                if _inline_plain_text(child.get("content", [])).strip():
                    yield {"kind": "paragraph", "level": list_level + 1, "content": child.get("content", [])}
                continue
            if child_type == "bulletList":
                yield from _iter_list_elements(child.get("content", []), ordered=False, list_level=list_level + 1)
                continue
            if child_type == "orderedList":
                child_start = max(int(child.get("attrs", {}).get("start", 1) or 1), 1)
                yield from _iter_list_elements(
                    child.get("content", []),
                    ordered=True,
                    list_level=list_level + 1,
                    start=child_start,
                )
        if ordered:
            index += 1


def render_markdown_bytes(content: str) -> bytes:
    return content.encode("utf-8")


def render_docx_bytes(content: str) -> bytes:
    document = Document()
    blocks = markdown_to_blocks(content)
    for element in _iter_renderable_elements(blocks):
        if element["kind"] == "heading":
            paragraph = document.add_heading(level=min(element["level"], 4))
            paragraph.text = ""
            _append_docx_runs(paragraph, element["content"])
            continue

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.3 * element["level"])
        if element["kind"] == "list_item":
            marker = f"{element['index']}." if element["ordered"] else "-"
            paragraph.add_run(f"{marker} ")
        _append_docx_runs(paragraph, element["content"])

    if not document.paragraphs:
        document.add_paragraph("")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_pdf_bytes(content: str) -> bytes:
    buffer = io.BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54,
    )
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "ExportBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=14,
        spaceAfter=6,
    )
    heading_styles = {
        1: ParagraphStyle("ExportHeading1", parent=styles["Heading1"], spaceAfter=10),
        2: ParagraphStyle("ExportHeading2", parent=styles["Heading2"], spaceAfter=8),
        3: ParagraphStyle("ExportHeading3", parent=styles["Heading3"], spaceAfter=6),
        4: ParagraphStyle("ExportHeading4", parent=styles["Heading4"], spaceAfter=6),
    }

    flowables = []
    blocks = markdown_to_blocks(content)
    for element in _iter_renderable_elements(blocks):
        if element["kind"] == "heading":
            style = heading_styles.get(element["level"], heading_styles[4])
            flowables.append(Paragraph(_inline_reportlab_markup(element["content"]) or "&nbsp;", style))
            flowables.append(Spacer(1, 2))
            continue

        style = ParagraphStyle(
            f"ExportBodyLevel{element['level']}_{element['kind']}",
            parent=body_style,
            leftIndent=18 * element["level"],
        )
        markup = _inline_reportlab_markup(element["content"]) or "&nbsp;"
        if element["kind"] == "list_item":
            marker = f"{element['index']}." if element["ordered"] else "&bull;"
            markup = f"{marker} {markup}"
        flowables.append(Paragraph(markup, style))

    if not flowables:
        flowables.append(Paragraph("&nbsp;", body_style))
    document.build(flowables)
    return buffer.getvalue()


def render_export_bytes(artifact: ExportArtifact) -> tuple[bytes, str]:
    file_type = export_file_type_for_artifact(artifact)
    if file_type == ExportFileType.PDF:
        payload = render_pdf_bytes(artifact.content)
    elif file_type == ExportFileType.DOCX:
        payload = render_docx_bytes(artifact.content)
    else:
        payload = render_markdown_bytes(artifact.content)
    return payload, export_content_type(file_type)


def toggle_share(artifact, enabled: bool):
    artifact.share_enabled = enabled
    if enabled and not artifact.share_token:
        artifact.share_token = secrets.token_hex(12)
    artifact.save(update_fields=["share_enabled", "share_token", "updated_at"])
    return artifact
