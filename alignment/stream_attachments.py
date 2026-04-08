from __future__ import annotations

import json
import os
from dataclasses import dataclass

from django.db import transaction
from django.utils import timezone
from docx import Document
from pypdf import PdfReader

from alignment.models import StreamAttachment, StreamAttachmentExtractionStatus, StreamPost, StreamPostKind
from specs.models import AIUsageOperation
from specs.openai import OpenAIUsageContext, request_openai_json_schema
from specs.services import apply_batch_spec_operations, section_summaries

SUPPORTED_STREAM_ATTACHMENT_EXTENSIONS = {
    ".txt": "text",
    ".md": "markdown",
    ".markdown": "markdown",
    ".pdf": "pdf",
    ".docx": "docx",
}
MAX_STREAM_ATTACHMENTS_PER_POST = 5
MAX_STREAM_ATTACHMENT_SIZE_BYTES = 20 * 1024 * 1024
STREAM_ATTACHMENT_RAW_TEXT_BUDGET = 18_000
STREAM_ATTACHMENT_SUMMARY_INPUT_CHAR_BUDGET = 12_000
STREAM_ATTACHMENT_AGENT_NAME = "SpecBridge AI"
STREAM_ATTACHMENT_AGENT_TITLE = "Agent"

ATTACHMENT_SUMMARY_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "properties": {
        "summary": {"type": "string", "minLength": 1},
        "key_points": {
            "type": "array",
            "items": {"type": "string", "minLength": 1},
        },
    },
    "required": ["summary", "key_points"],
}

STREAM_SPEC_APPLY_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "properties": {
        "summary": {"type": "string", "minLength": 1},
        "operations": {
            "type": "array",
            "items": {
                "anyOf": [
                    {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "type": {"type": "string", "const": "update_section"},
                            "section_id": {"type": "string", "minLength": 1},
                            "body": {"type": "string", "minLength": 1},
                        },
                        "required": ["type", "section_id", "body"],
                    },
                    {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "type": {"type": "string", "const": "insert_section_after"},
                            "after_section_id": {"type": "string", "minLength": 1},
                            "title": {"type": "string", "minLength": 1},
                            "body": {"type": "string", "minLength": 1},
                        },
                        "required": ["type", "after_section_id", "title", "body"],
                    },
                ]
            },
        },
    },
    "required": ["summary", "operations"],
}


class StreamAttachmentValidationError(Exception):
    def __init__(self, errors: dict[str, list[str]]):
        self.errors = errors
        super().__init__("Invalid stream attachment upload.")


class StreamAttachmentExtractionError(Exception):
    pass


class StreamSpecApplyError(Exception):
    pass


@dataclass
class StreamSpecApplyResult:
    summary: str
    applied_operations: list[dict]
    project_revision: object
    agent_post: StreamPost


def touch_project_activity(project) -> None:
    project.last_activity_at = timezone.now()
    project.save(update_fields=["last_activity_at", "updated_at"])


def normalize_attachment_extension(filename: str) -> str:
    return os.path.splitext(filename or "")[1].strip().lower()


def attachment_download_url(attachment: StreamAttachment) -> str:
    return f"/api/projects/{attachment.project.slug}/files/{attachment.id}/download"


def serialize_stream_attachment(attachment: StreamAttachment) -> dict:
    return {
        "id": attachment.id,
        "original_name": attachment.original_name,
        "content_type": attachment.content_type,
        "size_bytes": attachment.size_bytes,
        "extension": attachment.extension,
        "created_at": attachment.created_at.isoformat(),
        "extracted_char_count": attachment.extracted_char_count,
        "extraction_status": attachment.extraction_status,
        "extraction_error": attachment.extraction_error,
        "download_url": attachment_download_url(attachment),
    }


def validate_stream_uploads(uploaded_files) -> None:
    errors: list[str] = []
    files = list(uploaded_files or [])
    if len(files) > MAX_STREAM_ATTACHMENTS_PER_POST:
        errors.append(f"Upload at most {MAX_STREAM_ATTACHMENTS_PER_POST} files per post.")
    for uploaded_file in files:
        extension = normalize_attachment_extension(getattr(uploaded_file, "name", ""))
        if extension not in SUPPORTED_STREAM_ATTACHMENT_EXTENSIONS:
            errors.append(
                f"{getattr(uploaded_file, 'name', 'File')}: unsupported type. "
                "Use TXT, MD, PDF, or DOCX."
            )
        if getattr(uploaded_file, "size", 0) > MAX_STREAM_ATTACHMENT_SIZE_BYTES:
            errors.append(
                f"{getattr(uploaded_file, 'name', 'File')}: file exceeds the 20 MB limit."
            )
    if errors:
        raise StreamAttachmentValidationError({"files": errors})


def _read_text_file(file_obj) -> str:
    raw_bytes = file_obj.read()
    if not raw_bytes:
        raise StreamAttachmentExtractionError("The uploaded text file is empty.")
    return raw_bytes.decode("utf-8-sig", errors="replace").strip()


def _read_pdf_file(file_obj) -> str:
    reader = PdfReader(file_obj)
    parts = [(page.extract_text() or "").strip() for page in reader.pages]
    text = "\n\n".join(part for part in parts if part).strip()
    if not text:
        raise StreamAttachmentExtractionError("No extractable text was found in the PDF.")
    return text


def _read_docx_file(file_obj) -> str:
    document = Document(file_obj)
    parts: list[str] = []
    parts.extend(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()).strip()
            if row_text:
                parts.append(row_text)
    text = "\n\n".join(part for part in parts if part).strip()
    if not text:
        raise StreamAttachmentExtractionError("No extractable text was found in the DOCX file.")
    return text


def extract_text_from_attachment(attachment: StreamAttachment) -> str:
    extension = attachment.extension
    with attachment.stored_file.open("rb") as file_obj:
        if extension in {".txt", ".md", ".markdown"}:
            return _read_text_file(file_obj)
        if extension == ".pdf":
            return _read_pdf_file(file_obj)
        if extension == ".docx":
            return _read_docx_file(file_obj)
    raise StreamAttachmentExtractionError("Unsupported file type.")


@transaction.atomic
def attach_files_to_post(post: StreamPost, uploaded_files) -> list[StreamAttachment]:
    attachments: list[StreamAttachment] = []
    for uploaded_file in uploaded_files or []:
        extension = normalize_attachment_extension(getattr(uploaded_file, "name", ""))
        attachment = StreamAttachment.objects.create(
            project=post.project,
            post=post,
            stored_file=uploaded_file,
            original_name=getattr(uploaded_file, "name", "upload"),
            content_type=getattr(uploaded_file, "content_type", "") or "",
            size_bytes=max(int(getattr(uploaded_file, "size", 0) or 0), 0),
            extension=extension,
        )
        try:
            extracted_text = extract_text_from_attachment(attachment)
            attachment.extracted_text = extracted_text
            attachment.extracted_char_count = len(extracted_text)
            attachment.extraction_status = StreamAttachmentExtractionStatus.COMPLETED
            attachment.extraction_error = ""
        except StreamAttachmentExtractionError as exc:
            attachment.extracted_text = ""
            attachment.extracted_char_count = 0
            attachment.extraction_status = StreamAttachmentExtractionStatus.FAILED
            attachment.extraction_error = str(exc)
        attachment.save(
            update_fields=[
                "extracted_text",
                "extracted_char_count",
                "extraction_status",
                "extraction_error",
                "updated_at",
            ]
        )
        attachments.append(attachment)
    return attachments


def _request_openai(
    *,
    schema_name: str,
    schema: dict,
    prompt: str,
    project,
    actor=None,
    operation: str,
    context_metadata: dict | None = None,
) -> tuple[str, dict]:
    response = request_openai_json_schema(
        schema_name=schema_name,
        schema=schema,
        prompt=prompt,
        error_cls=StreamSpecApplyError,
        empty_output_message="OpenAI returned an empty attachment response.",
        usage_context=OpenAIUsageContext(
            project=project,
            actor=actor,
            operation=operation,
            context_metadata=context_metadata or {},
        ),
    )
    try:
        return response.model, json.loads(response.output_text)
    except json.JSONDecodeError as exc:
        raise StreamSpecApplyError(f"OpenAI returned malformed JSON: {response.output_text}") from exc


def _trim_text_for_summary(text: str, budget: int = STREAM_ATTACHMENT_SUMMARY_INPUT_CHAR_BUDGET) -> str:
    normalized = (text or "").strip()
    if len(normalized) <= budget:
        return normalized
    head_budget = budget // 2
    tail_budget = budget - head_budget
    return f"{normalized[:head_budget].rstrip()}\n\n[...]\n\n{normalized[-tail_budget:].lstrip()}"


def _summary_prompt(*, attachment: StreamAttachment, text: str) -> str:
    return (
        "You are summarizing an uploaded reference document for a collaborative product spec workspace.\n"
        "Return a concise, implementation-relevant summary in English.\n"
        "Call out requirements, constraints, integrations, APIs, compliance notes, and non-obvious caveats when present.\n"
        "Do not invent content that is not supported by the document.\n"
        f"File name: {attachment.original_name}\n"
        f"File type: {attachment.extension or attachment.content_type}\n\n"
        "Document text:\n"
        f"{text}"
    )


def summarize_attachment_for_apply(attachment: StreamAttachment, *, actor=None) -> dict:
    _, parsed = _request_openai(
        schema_name="stream_attachment_summary",
        schema=ATTACHMENT_SUMMARY_SCHEMA,
        prompt=_summary_prompt(
            attachment=attachment,
            text=_trim_text_for_summary(attachment.extracted_text),
        ),
        project=attachment.project,
        actor=actor,
        operation=AIUsageOperation.ATTACHMENT_SUMMARY,
        context_metadata={"attachment_id": attachment.id, "extension": attachment.extension},
    )
    return {
        "id": attachment.id,
        "name": attachment.original_name,
        "type": attachment.extension or attachment.content_type,
        "mode": "summary",
        "summary": parsed["summary"].strip(),
        "key_points": [item.strip() for item in parsed.get("key_points", []) if item.strip()],
    }


def _docs_payload_for_apply(attachments: list[StreamAttachment], *, actor=None) -> list[dict]:
    total_chars = sum(len(attachment.extracted_text or "") for attachment in attachments)
    if total_chars <= STREAM_ATTACHMENT_RAW_TEXT_BUDGET:
        return [
            {
                "id": attachment.id,
                "name": attachment.original_name,
                "type": attachment.extension or attachment.content_type,
                "mode": "raw_text",
                "text": attachment.extracted_text,
            }
            for attachment in attachments
        ]
    return [summarize_attachment_for_apply(attachment, actor=actor) for attachment in attachments]


def _stream_apply_prompt(*, prompt: str, sections: list[dict], documents: list[dict]) -> str:
    section_payload = [
        {
            "id": section["id"],
            "title": section["title"],
            "status": section["status"],
            "kind": section["kind"],
            "body": section["body"],
        }
        for section in sections
    ]
    return (
        "You are applying uploaded reference documents to a collaborative product specification.\n"
        "Interpret the user's prompt as a direct instruction to mutate the current spec.\n"
        "Return spec-ready English content only.\n"
        "You may do two things only:\n"
        "1. update_section: replace the body of an existing section.\n"
        "2. insert_section_after: add a new custom section after an existing section.\n"
        "Do not rename, delete, or reorder existing sections.\n"
        "Do not invent requirements that are unsupported by the uploaded documents and current spec.\n"
        "Prefer editing existing sections when possible, and insert a new section only when the request clearly needs one.\n"
        f"User prompt:\n{prompt}\n\n"
        f"Current sections JSON:\n{json.dumps(section_payload, ensure_ascii=True)}\n\n"
        f"Uploaded documents JSON:\n{json.dumps(documents, ensure_ascii=True)}"
    )


def _normalized_apply_operations(parsed_output: dict, sections: list[dict]) -> tuple[str, list[dict]]:
    section_ids = {section["id"] for section in sections}
    operations: list[dict] = []
    for operation in parsed_output.get("operations", []):
        operation_type = (operation.get("type") or "").strip()
        if operation_type == "update_section":
            section_id = (operation.get("section_id") or "").strip()
            body = (operation.get("body") or "").strip()
            if section_id in section_ids and body:
                operations.append({"type": operation_type, "section_id": section_id, "body": body})
        elif operation_type == "insert_section_after":
            after_section_id = (operation.get("after_section_id") or "").strip()
            title = (operation.get("title") or "").strip()
            body = (operation.get("body") or "").strip()
            if after_section_id in section_ids and title and body:
                operations.append(
                    {
                        "type": operation_type,
                        "after_section_id": after_section_id,
                        "title": title,
                        "body": body,
                    }
                )
    return (parsed_output.get("summary") or "").strip(), operations


def _success_notice_body(summary: str, operations: list[dict]) -> str:
    lines = [summary.strip() or "Applied the uploaded document guidance to the spec."]
    for operation in operations:
        if operation["type"] == "update_section":
            lines.append(f"- Updated {operation['section_title']}")
        elif operation["type"] == "insert_section_after":
            lines.append(f"- Added {operation['section_title']}")
    return "\n".join(lines)


def process_uploaded_documents_for_post(post: StreamPost, *, prompt: str, actor=None) -> StreamSpecApplyResult:
    attachments = list(post.attachments.all())
    failed_attachments = [
        attachment.original_name
        for attachment in attachments
        if attachment.extraction_status != StreamAttachmentExtractionStatus.COMPLETED
    ]
    if failed_attachments:
        raise StreamSpecApplyError(
            "Couldn't extract usable text from: " + ", ".join(failed_attachments) + "."
        )

    sections = section_summaries(post.project)
    if not sections:
        raise StreamSpecApplyError("The current spec has no sections to update.")

    _, parsed_output = _request_openai(
        schema_name="stream_spec_apply",
        schema=STREAM_SPEC_APPLY_SCHEMA,
        prompt=_stream_apply_prompt(
            prompt=prompt,
            sections=sections,
            documents=_docs_payload_for_apply(attachments, actor=actor),
        ),
        project=post.project,
        actor=actor,
        operation=AIUsageOperation.STREAM_SPEC_APPLY,
        context_metadata={"post_id": post.id, "attachment_ids": [attachment.id for attachment in attachments]},
    )
    summary, operations = _normalized_apply_operations(parsed_output, sections)
    if not operations:
        raise StreamSpecApplyError("AI did not return any concrete spec changes to apply.")

    apply_result = apply_batch_spec_operations(
        project=post.project,
        operations=operations,
        actor=actor,
        title="Applied uploaded document guidance",
        summary=summary or post.body.strip() or "Applied uploaded document guidance.",
        source_post=post,
    )
    if not apply_result["project_revision"]:
        raise StreamSpecApplyError("No spec changes were applied.")

    agent_post = StreamPost.objects.create(
        project=post.project,
        actor_name=STREAM_ATTACHMENT_AGENT_NAME,
        actor_title=STREAM_ATTACHMENT_AGENT_TITLE,
        kind=StreamPostKind.AGENT,
        concern=post.concern,
        body=_success_notice_body(summary, apply_result["applied_operations"]),
    )
    touch_project_activity(post.project)
    return StreamSpecApplyResult(
        summary=summary,
        applied_operations=apply_result["applied_operations"],
        project_revision=apply_result["project_revision"],
        agent_post=agent_post,
    )


def create_failed_upload_notice(post: StreamPost, *, message: str) -> StreamPost:
    agent_post = StreamPost.objects.create(
        project=post.project,
        actor_name=STREAM_ATTACHMENT_AGENT_NAME,
        actor_title=STREAM_ATTACHMENT_AGENT_TITLE,
        kind=StreamPostKind.AGENT,
        concern=post.concern,
        body=(
            "I couldn't apply the uploaded files to the spec. "
            f"{message.strip()} The files are still attached to your post, and the spec was not changed."
        ),
    )
    touch_project_activity(post.project)
    return agent_post
